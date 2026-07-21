# -*- coding: utf-8 -*-
"""
신한 템플릿 기반 DOCX 생성 스크립트
server.mjs에서 호출됨.

사용법:
  python scripts/build_docx.py --type sokbo --company "현대로템" --date "2026년 7월 21일" --output output.docx < markdown.txt
"""
import sys, os, io, argparse, re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn as docx_qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE_DIR = Path(r"C:\Users\lucy0\Desktop\신한_기본템플릿")


def _find_sect_break_paras(body_elem):
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    result = []
    for p in body_elem.iterchildren(f'{ns}p'):
        pPr = p.find(f'{ns}pPr')
        if pPr is not None and pPr.find(f'{ns}sectPr') is not None:
            result.append(p)
    return result


def _parse_md(md_text):
    elements = []
    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith('|') and s.count('|') >= 3:
            tbl = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl.append(lines[i].strip())
                i += 1
            elements.append(('table', tbl))
            continue
        if s.startswith('### '):
            elements.append(('Title1', s[4:]))
        elif s.startswith('## '):
            elements.append(('Title2', s[3:]))
        elif s.startswith('# '):
            elements.append(('Title3_라인', s[2:]))
        elif s.startswith('■'):
            elements.append(('Title3_라인', s))
        elif s.startswith('- ') or s.startswith('▸ '):
            elements.append(('bullet', s[2:]))
        elif s == '주요 Q&A':
            elements.append(('Title3_라인', s))
        elif s.startswith('Q.') or s.startswith('Q '):
            elements.append(('Title1', s))
        else:
            elements.append(('본문1', s))
        i += 1
    return elements


def _add_styled_para(doc, body, text, style_name, insert_ref=None):
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()
    parts = text.split('**')
    for idx, part in enumerate(parts):
        if part:
            run = p.add_run(part)
            if idx % 2 == 1:
                run.bold = True
    if insert_ref is not None:
        body.remove(p._element)
        insert_ref.addprevious(p._element)
    return p


def _add_md_table(doc, body, table_lines, insert_ref=None):
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip('|').split('|')]
        if cells and all(set(c) <= {'-', ':', ' ', ''} for c in cells):
            continue
        if cells:
            rows_data.append(cells)
    if not rows_data:
        return
    ncols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=ncols)
    for ri, rd in enumerate(rows_data):
        for ci, ct in enumerate(rd):
            if ci < ncols:
                cell = tbl.rows[ri].cells[ci]
                cell.paragraphs[0].text = ''
                try:
                    cell.paragraphs[0].style = doc.styles['tableCB' if ri == 0 else 'tableL']
                except KeyError:
                    pass
                cell.paragraphs[0].add_run(ct)
    if insert_ref is not None:
        tbl_elem = tbl._tbl
        body.remove(tbl_elem)
        insert_ref.addprevious(tbl_elem)


def _fill_sokbo_table_cell(doc, elements):
    tbl = doc.tables[0]
    cell = tbl.rows[9].cells[1]

    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)

    sokbo_style_map = {
        'Title2': 'Title2',
        'Title1': 'Title1',
        '본문1': '속보_본문1',
        'Title3_라인': 'Title2',
        'bullet': '속보_본문1',
    }

    _skip_patterns = [
        re.compile(r'^신한\s*속보'),
        re.compile(r'^\[?\s*Analyst\s*\]?'),
        re.compile(r'^주가\s*[\(（]'),
        re.compile(r'^시가총액'),
        re.compile(r'^\d{4}년\s*\d{1,2}월\s*\d{1,2}일'),
        re.compile(r'^이동헌\s*연구위원'),
        re.compile(r'^이지한\s*연구원'),
    ]

    def _is_header_line(text):
        t = text.strip()
        for pat in _skip_patterns:
            if pat.search(t):
                return True
        return False

    for etype, content in elements:
        if etype == 'table':
            continue
        if _is_header_line(content):
            continue
        content = _clean_text(content.replace('—', ', ').replace('–', ', '))
        style_name = sokbo_style_map.get(etype, '속보_본문1')
        p = cell.add_paragraph()
        try:
            p.style = doc.styles[style_name]
        except KeyError:
            pass
        parts = content.split('**')
        for idx, part in enumerate(parts):
            if part:
                run = p.add_run(_clean_text(part))
                if idx % 2 == 1:
                    run.bold = True


def build_note_docx(raw_md, company, pub_date_kr):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.3

    cleaned = '\n'.join(
        line for line in raw_md.split('\n')
        if not line.strip().startswith('발간일:') and not line.strip().startswith('기업명:') and not line.strip().startswith('분기:') and line.strip() != '---'
    )

    for line in cleaned.split('\n'):
        s = line.strip()
        if not s:
            continue
        if s.startswith('**') and s.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(s[2:-2])
            run.bold = True
            run.font.size = Pt(10)
            continue
        if s.startswith('Q)') or s.startswith('Q '):
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.bold = True
            run.font.size = Pt(10)
            continue
        p = doc.add_paragraph()
        parts = s.split('**')
        for idx, part in enumerate(parts):
            if part:
                run = p.add_run(part)
                run.font.size = Pt(10)
                if idx % 2 == 1:
                    run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_docx_from_template(raw_md, report_type, company, pub_date_kr):
    if report_type == 'note':
        return build_note_docx(raw_md, company, pub_date_kr)

    is_sokbo = (report_type == 'sokbo')
    tpl = TEMPLATE_DIR / ("신한 속보_기본템플릿.docx" if is_sokbo else "신한 레포트_기본템플릿.docx")

    if not tpl.exists():
        print(f"[경고] 템플릿 파일 없음: {tpl}. 기본 DOCX로 생성합니다.", file=sys.stderr)
        return build_note_docx(raw_md, company, pub_date_kr)

    doc = Document(str(tpl))
    body = doc.element.body
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # 머리글 날짜 + 기업명 교체
    wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for sec in doc.sections:
        if not sec.header:
            continue
        all_t = list(sec.header._element.iter(f'{{{wns}}}t'))
        date_start = None
        for i, t in enumerate(all_t):
            txt = t.text or ''
            if '년' in txt:
                date_start = i
                for j in range(i - 1, max(i - 3, -1), -1):
                    if any(c.isdigit() for c in (all_t[j].text or '')):
                        date_start = j
                    else:
                        break
            if date_start is not None and '일' in txt:
                all_t[date_start].text = pub_date_kr
                for j in range(date_start + 1, i + 1):
                    all_t[j].text = ''
                date_start = None
        if company:
            for t in all_t:
                if t.text in ('기업명', '회사명', '섹터명'):
                    t.text = company
            for idx in range(len(all_t) - 1):
                if (all_t[idx].text or '') == '회' and (all_t[idx + 1].text or '') == '사명':
                    all_t[idx].text = company
                    all_t[idx + 1].text = ''

    # 속보 표지 테이블 날짜 + 섹터명
    if is_sokbo and doc.tables:
        for cell in doc.tables[0].rows[0].cells:
            for p in cell.paragraphs:
                if 'X월' in p.text or ('2026' in p.text and '월' in p.text):
                    for run in p.runs:
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = pub_date_kr
                    else:
                        p.add_run(pub_date_kr)
        if company:
            for cell in doc.tables[0].rows[3].cells:
                for p in cell.paragraphs:
                    if '섹터명' in p.text:
                        for run in p.runs:
                            if '섹터명' in (run.text or ''):
                                run.text = run.text.replace('섹터명', company)
                        break

    # 발간일/기업명/분기 줄 제거 후 파싱
    cleaned_md = '\n'.join(
        line for line in raw_md.split('\n')
        if not line.strip().startswith('발간일:') and not line.strip().startswith('기업명:') and not line.strip().startswith('분기:') and line.strip() != '---'
    )

    elements = _parse_md(cleaned_md)

    if is_sokbo and doc.tables:
        _fill_sokbo_table_cell(doc, elements)
    else:
        sect_breaks = _find_sect_break_paras(body)
        insert_ref = None
        keep = 0
        if sect_breaks:
            all_p = [c for c in body.iterchildren(f'{ns}p')]
            keep = all_p.index(sect_breaks[0]) + 1
        for p in list(doc.paragraphs[keep:]):
            body.remove(p._element)

        for etype, content in elements:
            if etype == 'table':
                _add_md_table(doc, body, content, insert_ref)
            else:
                _add_styled_para(doc, body, content, etype, insert_ref)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _clean_text(text):
    """서로게이트 문자 및 XML 호환 불가 문자 제거."""
    import re
    # 서로게이트 페어 제거
    text = text.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
    # XML 1.0 허용 불가 제어문자 제거
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--type', default='qa', choices=['qa', 'sokbo', 'review', 'overseas', 'note'])
    parser.add_argument('--company', default='레포트')
    parser.add_argument('--date', default='')
    parser.add_argument('--output', required=True)
    args = parser.parse_args()

    # stdin을 surrogateescape로 읽고 클린업
    sys.stdin.reconfigure(encoding='utf-8', errors='surrogateescape')
    raw_md = _clean_text(sys.stdin.read())
    if not raw_md.strip():
        print("마크다운 입력이 비어있습니다.", file=sys.stderr)
        sys.exit(1)

    buf = build_docx_from_template(raw_md, args.type, args.company, args.date)

    with open(args.output, 'wb') as f:
        f.write(buf.read())

    print(f"[완료] {args.output}", file=sys.stderr)


if __name__ == '__main__':
    main()
